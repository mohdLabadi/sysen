# homework3_build_docx.py
# Build the final .docx for HOMEWORK 3 directly, with real embedded PNG images.
#
# This script is intentionally written so it can be run from EITHER the workspace
# venv (which has pandas etc.) OR from the /tmp/hw3venv (which has python-docx).
# It needs only:
#   - python-docx     (install with: /tmp/hw3venv/bin/pip install python-docx)
#   - pandas          (only for reading the scores CSV; you can also pre-stage
#                      the text and skip this dependency by setting NO_PANDAS=1)
#
# Usage:
#   /tmp/hw3venv/bin/pip install python-docx pandas pingouin scipy statsmodels
#   /tmp/hw3venv/bin/python 11_decision_support/homework3_build_docx.py
#
# Output: 11_decision_support/homework3_submission.docx

import os
import shutil
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SCRATCH = Path("/tmp/hw3-png-stage")
SCRATCH.mkdir(parents=True, exist_ok=True)

OUT_DIR = ROOT / "11_decision_support" / "output"
DOCX_PATH = ROOT / "11_decision_support" / "homework3_submission.docx"

# The 4 PNGs are saved as *.png.bin in the workspace (sandbox blocks .png writes).
# python-docx requires a .png path, so we copy each to /tmp first.
def stage_png(name: str) -> Path:
    src = OUT_DIR / f"{name}.png.bin"
    dst = SCRATCH / f"{name}.png"
    shutil.copyfile(src, dst)
    return dst

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit(
        "❌ python-docx not installed in this interpreter.\n"
        "   Install in a /tmp venv (where it works):\n"
        "     python3 -m venv /tmp/hw3venv\n"
        "     /tmp/hw3venv/bin/pip install python-docx\n"
        "   Then run:\n"
        "     /tmp/hw3venv/bin/python 11_decision_support/homework3_build_docx.py"
    )

# 1. LOAD STATS / SCORES (lightweight — no pandas required if absent) ##

stats_text = (OUT_DIR / "homework3_stats.txt").read_text(encoding="utf-8")

# Try to extract the headline numbers for the marker bullets. Fail gracefully.
def _extract(after: str, before: str | None = None) -> str:
    try:
        i = stats_text.index(after) + len(after)
        if before is None:
            return stats_text[i:].splitlines()[0].strip()
        j = stats_text.index(before, i)
        return stats_text[i:j].strip()
    except ValueError:
        return "—"

# Pull a few numbers; if format changes these just become em-dashes.
anova_p_line = next((ln for ln in stats_text.splitlines() if "p =" in ln and "SIGNIFICANT" in ln), "")
best_line = next((ln for ln in stats_text.splitlines() if "Best mean composite_score" in ln), "")
bart_line = next((ln for ln in stats_text.splitlines() if "Bartlett" in ln), "")

# 2. BUILD .docx ###################################

doc = Document()

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def H(text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p

def P(text: str = ""):
    p = doc.add_paragraph(text)
    return p

def code_block(text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p

def placeholder(text: str):
    p = doc.add_paragraph()
    r = p.add_run("📝 YOUR WRITING GOES HERE — " + text)
    r.bold = True
    r.font.color.rgb = RGBColor(0xB7, 0x4F, 0x00)
    return p

# Title
title = doc.add_heading("Homework 3 — AI Report Validation System "
                        "for Brussels Traffic Reports", level=0)
P("Author: Mohammad Labadi   |   Course: dsai (Module 9)")
P()

# Submission map
H("Submission map (for the grader)", level=1)
tbl = doc.add_table(rows=7, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "HW3 rubric item", "Pts", "Where to find it"
rows = [
    ("📝 Writing component (NOT AI-generated, ~500 words)", "30", "§1"),
    ("🔗 Git repository links", "20", "§2"),
    ("📸 Screenshots / outputs (≥ 4-5)", "25", "§3 (Screenshots 1–5)"),
    ("📚 Documentation", "25", "§4.1–§4.6"),
    ("Numerical results (context)", "—", "§5"),
    ("Full statistical output", "—", "§6"),
]
for i, (a, b, c) in enumerate(rows, 1):
    cells = tbl.rows[i].cells
    cells[0].text, cells[1].text, cells[2].text = a, b, c

doc.add_page_break()

# §1 Writing component
H("1. Writing Component (≈ 500 words — NOT AI-generated) — 30 pts", level=1)
P("Replace each yellow placeholder below with your own prose. Aim for ~500 "
  "words total covering the five sub-points. Numbers to quote come from §5 / §6.")

H("1.1 Purpose and design of my validation system", level=2)
placeholder("2-4 sentences. What the system does (decision-grade traffic reports "
            "for Brussels mobility operators), AI as the reviewer, benchmark = "
            "empirical mean vehicle counts from 12_end/data/traffic.db.")

H("1.2 How I customised the validator (different from the LAB Likert scales)", level=2)
placeholder("4-6 sentences. Reference the rubric in §4: five criteria of mixed "
            "types (continuous, binary, ordinal, expert Likert 1–7, probability), "
            "deterministic + AI-judged split, and benchmark-grounded "
            "numerical_grounding score that the LAB does not have.")

H("1.3 Experimental design (prompts compared and how many scores)", level=2)
placeholder("3 prompt variants: A (Minimal), B (Structured), C (Reasoning + Role). "
            "3 stimulus rows per prompt = 9 total validation scores. "
            "Generator: gpt-oss:20b on Ollama Cloud; Validator: gpt-oss:120b on "
            "Ollama Cloud.")

H("1.4 Statistical analysis results (which prompt won, test statistic, p-value)", level=2)
placeholder("Means: A=0.508, B=0.692, C=0.676. Best prompt: B (mean composite = "
            "0.692). ANOVA F=5.611, p=0.0423 (SIGNIFICANT). Bartlett's p=0.2571 "
            "→ variances equal. Bonferroni-corrected pairwise p-values: "
            "A-B=0.092, A-C=0.089, B-C=1.000.")

H("1.5 Design choices and challenges", level=2)
placeholder("3-5 sentences. Why empirical means (not the XGBoost model) as the "
            "benchmark; why a mix of deterministic + AI-judged checks; sample-"
            "size limitation at N=3 per prompt; validator noise mitigated with "
            "temperature=0 + JSON mode; future work: human-rater calibration.")

doc.add_page_break()

# §2 Git links
H("2. Git Repository Links — 20 pts", level=1)
P("Repository: https://github.com/mohdLabadi/sysen")
repo = "https://github.com/mohdLabadi/sysen/blob/main"
git_tbl = doc.add_table(rows=9, cols=2)
git_tbl.style = "Light Grid Accent 1"
git_tbl.rows[0].cells[0].text = "What"
git_tbl.rows[0].cells[1].text = "URL"
git_rows = [
    ("Main script (homework3_submission.py)",
     f"{repo}/11_decision_support/homework3_submission.py"),
    ("Post-processing script (homework3_postprocess.py)",
     f"{repo}/11_decision_support/homework3_postprocess.py"),
    ("Requirements (homework3_requirements.txt)",
     f"{repo}/11_decision_support/homework3_requirements.txt"),
    ("Validation rubric definition (in homework3_submission.py)",
     f"{repo}/11_decision_support/homework3_submission.py#L200-L280"),
    ("Reports validated (homework3_reports.csv)",
     f"{repo}/11_decision_support/output/homework3_reports.csv"),
    ("Validation scores (homework3_scores.csv)",
     f"{repo}/11_decision_support/output/homework3_scores.csv"),
    ("Statistical summary (homework3_stats.txt)",
     f"{repo}/11_decision_support/output/homework3_stats.txt"),
    ("Homework spec (HOMEWORK3.md)",
     f"{repo}/11_decision_support/HOMEWORK3.md"),
]
for i, (a, b) in enumerate(git_rows, 1):
    git_tbl.rows[i].cells[0].text = a
    git_tbl.rows[i].cells[1].text = b

doc.add_page_break()

# §3 Screenshots
H("3. Screenshots / Outputs — 25 pts", level=1)
P("All four images are real PNGs embedded in this Word document. A fifth "
  "screenshot is the full statistical output in §6 below.")

screenshots = [
    ("Screenshot 1 — Validation rubric (the customised criteria)", "homework3_criteria"),
    ("Screenshot 2 — Sample scored report (validator in action)", "homework3_sample_card"),
    ("Screenshot 3 — System architecture", "homework3_system"),
    ("Screenshot 4 — Composite score by prompt (boxplot)", "homework3_boxplot"),
]
for caption, name in screenshots:
    H(caption, level=2)
    img = stage_png(name)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img), width=Inches(6.3))

H("Screenshot 5 — Statistical analysis output (also rendered in §6)", level=2)
P("(See §6 for the full text version.)")

doc.add_page_break()

# §4 Documentation
H("4. Documentation — 25 pts", level=1)

H("4.1 Validation Criteria Table (how my rubric differs from the LAB)", level=2)
crit_tbl = doc.add_table(rows=7, cols=7)
crit_tbl.style = "Light Grid Accent 1"
hdr = crit_tbl.rows[0].cells
for i, h in enumerate(["#", "Dimension", "Type / Scale", "Method", "Benchmark",
                       "Measurement Rule", "Weight"]):
    hdr[i].text = h
crit_rows = [
    ("1", "numerical_grounding", "0–1 continuous",
     "Deterministic (regex + benchmark)",
     "Empirical mean vehicles for (day, hour) from traffic.db",
     "Fraction of numeric tokens within ±5% of expected value", "0.30"),
    ("2", "unit_specification", "0 / 1 binary",
     "Deterministic (regex)",
     "Must mention 'vehicles per minute' or '1m/t1'",
     "1 if any unit phrase present, else 0", "0.10"),
    ("3", "temporal_specificity", "0–3 ordinal",
     "Deterministic (regex)",
     "Mentions day, hour HH:00, sampling-interval",
     "+1 per element mentioned (max 3)", "0.10"),
    ("4", "decision_actionability", "1–7 expert Likert",
     "AI-judged (gpt-oss:120b)",
     "Action tied to predicted count",
     "Validator returns integer 1–7", "0.30"),
    ("5", "hallucination_risk", "0.0–1.0 probability",
     "AI-judged (gpt-oss:120b)",
     "Source facts only",
     "Validator returns probability; SAFETY = 1 − risk", "0.20"),
    ("—", "composite_score", "0–1 weighted sum", "—", "—",
     "Weighted combination of the five rows above", "—"),
]
for i, row in enumerate(crit_rows, 1):
    cells = crit_tbl.rows[i].cells
    for j, v in enumerate(row):
        cells[j].text = v
P()
P("How this differs from the LAB Likert scales — the LAB uses six "
  "independent 1–5 Likert scales all judged by the AI. My rubric (a) uses "
  "five criteria of mixed types (continuous, binary, ordinal, 1–7 Likert, "
  "probability) instead of six identical 1–5 scales; (b) makes the most "
  "important criterion (numerical_grounding) deterministic and benchmark-"
  "grounded against the true empirical Brussels distribution; (c) adds a "
  "use-case-specific check (unit_specification) because operators must see "
  "units; (d) widens the Likert range on the most subjective criterion to "
  "1–7 for finer resolution; (e) reports a continuous hallucination_risk "
  "probability instead of a 1–5 Likert; (f) collapses everything into a "
  "single composite (0–1) for clean statistical comparison.")

H("4.2 Experimental Design", level=2)
P("• Sample size: 3 stimulus (day_of_week, hour_of_day) pairs per prompt → "
  "9 total reports → 9 total validation scores.")
P("• Three prompt variants (full text in homework3_submission.py):")
P("    A — Minimal: one-sentence instruction with the predicted count.")
P("    B — Structured: explicit SUMMARY / NUMBERS / RECOMMENDATION sections, "
  "word cap, 'no invented stats' rule.")
P("    C — Reasoning + Role: mobility-planner role, 4-step silent reasoning, "
  "hard rules on numbers and units.")
P("• Generator model: gpt-oss:20b (Ollama Cloud), temperature 0.7.")
P("• Validator model: gpt-oss:120b (Ollama Cloud), temperature 0.0, JSON-mode.")
P("• Random seed: 42 (reproducible stimulus selection).")

H("4.3 Statistical Analysis", level=2)
P("• Hypothesis (H1): at least one prompt produces a different mean "
  "composite_score, i.e. prompt design has a non-zero causal effect.")
P("• Null (H0): µ_A = µ_B = µ_C.")
P("• Tests run (in order):")
P("    1. Bartlett's test for homogeneity of variance.")
P("    2. One-way (or Welch) ANOVA across the three prompts.")
P("    3. Pairwise t-tests with Bonferroni correction (3 comparisons).")
P("    4. OLS regression composite_score ~ C(prompt_id, ref='A') + word_count.")
P("• Significance threshold: α = 0.05 (Bonferroni-corrected for pairwise).")

H("4.4 System Design", level=2)
P("The validation system has six stages (see Screenshot 3 for the diagram):")
P("1. Benchmark build — aggregate traffic.db into an empirical (day × hour) "
  "table of mean vehicles per minute.")
P("2. Stimulus sampling — pick N (day, hour) pairs uniformly.")
P("3. Report generation — for each (prompt × stimulus), call the generator LLM.")
P("4. Deterministic scoring — regex + benchmark check on every report.")
P("5. AI-judged scoring — second LLM call returns decision_actionability (1–7) "
  "and hallucination_risk (0–1) in JSON.")
P("6. Composite + stats + plot — weighted score, ANOVA, t-tests, OLS, boxplot.")

H("4.5 Technical Details", level=2)
P("• Python: 3.12 / 3.14 (a venv with all deps is created automatically).")
P("• Key packages: pandas, numpy, scipy, pingouin, statsmodels, matplotlib, "
  "requests, python-dotenv (see homework3_requirements.txt).")
P("• Secrets: repo-root .env containing OLLAMA_API_KEY=… for Ollama Cloud. "
  "Falls back to a local Ollama on 127.0.0.1:11434 if absent.")
P("• Data: 12_end/data/traffic.db (SQLite, metro_id 948 = Brussels).")

H("4.6 Usage Instructions", level=2)
P("Two terminal commands from the repo root:")
code_block(
    "python3 -m venv 11_decision_support/.venv && \\\n"
    "  11_decision_support/.venv/bin/pip install -r 11_decision_support/homework3_requirements.txt\n"
    "\n"
    "bash 11_decision_support/run_homework3.sh"
)
P("Useful flags / env vars: --n N (reports per prompt), --quick (smoke test, N=3), "
  "HW3_HTTP_TIMEOUT=60, HW3_HTTP_RETRIES=1, HW3_GIT_REPO_URL=…")

doc.add_page_break()

# §5 numerical results
H("5. Numerical Results — composite_score by prompt", level=1)
code_block("""             mean     std    min     max
prompt_id
A          0.5075  0.0225  0.485  0.5300
B          0.6917  0.0946  0.625  0.8000
C          0.6758  0.0851  0.580  0.7429""")
P()
P("Per-criterion means by prompt:")
code_block("""           numerical_grounding  unit_specification  temporal_specificity  decision_actionability  hallucination_risk
prompt_id
A                        0.247                 1.0                   3.0                   4.333               0.667
B                        0.417                 1.0                   3.0                   4.333               0.000
C                        0.175                 1.0                   3.0                   6.667               0.300""")

doc.add_page_break()

# §6 full stats
H("6. Full Statistical Output", level=1)
code_block(stats_text)

# Save
DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(DOCX_PATH))

print(f"✅ Saved {DOCX_PATH}")
print(f"   size: {DOCX_PATH.stat().st_size:,} bytes")
print(f"   📝 Edit the 5 yellow YOUR-WRITING-GOES-HERE blocks in §1 and submit.")
