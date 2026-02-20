# 06_api_ai_reporter.py
# Query API + Process Data + AI Reporting Summary (Ollama)
# Pairs with 01_query_api/02_lab.py, 03_query_ai/02_ollama.py, and 03_query_ai/05_reporting.py
# Tim Fraser
#
# This script: queries Marketstack EOD data, cleans/aggregates it, sends a
# compact JSON payload to Ollama, and saves a single journalistic Word report.

# 0. Setup #################################

## 0.1 Load Packages ############################

import json
import os
from datetime import datetime

import pandas as pd
import requests
from dotenv import load_dotenv
from docx import Document

## 0.2 Configuration ############################

load_dotenv(".env")

MARKETSTACK_API_KEY = os.getenv("MARKETSTACK_API_KEY") or os.getenv("TEST_API_KEY2")
SYMBOL = os.getenv("MARKETSTACK_SYMBOL", "AAPL")
N_DAYS = int(os.getenv("MARKETSTACK_LIMIT", "60"))

OLLAMA_PORT = int(os.getenv("OLLAMA_PORT", "11434"))
OLLAMA_HOST = os.getenv("OLLAMA_HOST", f"http://localhost:{OLLAMA_PORT}")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "smollm2:1.7b")

OUTPUT_DOCX = f"03_query_ai/06_report_{SYMBOL.lower()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"


# 1. Query API #################################

def query_marketstack_eod(symbol: str, n_days: int, api_key: str) -> pd.DataFrame:
    """Query Marketstack EOD and return a DataFrame."""
    if not api_key:
        raise ValueError("Missing API key. Add MARKETSTACK_API_KEY to .env (project root).")
    url = "http://api.marketstack.com/v1/eod"
    params = {"access_key": api_key, "symbols": symbol, "limit": n_days, "sort": "DESC"}
    response = requests.get(url, params=params, timeout=30)
    print("API status code:", response.status_code)
    response.raise_for_status()
    records = response.json().get("data", [])
    if not records:
        raise ValueError("API returned 0 rows. Check symbol, key, and plan limits.")
    return pd.DataFrame(records)


# 2. Process Data #################################

def process_eod_data(df_raw: pd.DataFrame) -> tuple[pd.DataFrame, dict]:
    """Clean EOD data and compute reporting metrics. Returns (df_clean, metrics)."""
    df = df_raw.copy()
    keep_cols = [c for c in ["date", "open", "high", "low", "close", "volume", "symbol", "exchange"] if c in df.columns]
    df = df.filter(items=keep_cols)
    df["date"] = pd.to_datetime(df["date"], errors="coerce").dt.date
    df = df.dropna(subset=["date", "close"]).sort_values("date", ascending=True)
    for c in [x for x in ["open", "high", "low", "close", "volume"] if x in df.columns]:
        df[c] = pd.to_numeric(df[c], errors="coerce")
    df = df.dropna(subset=["close"])
    df = (df
          .assign(daily_return_pct=lambda x: x["close"].pct_change() * 100)
          .assign(rolling_vol_10d_pct=lambda x: x["daily_return_pct"].rolling(10).std()))
    start_close = float(df["close"].iloc[0])
    end_close = float(df["close"].iloc[-1])
    net_change_pct = (end_close / start_close - 1) * 100
    high_close = float(df["close"].max())
    low_close = float(df["close"].min())
    avg_volume = float(df["volume"].mean()) if "volume" in df.columns else None
    last_10 = df.tail(10)
    first_10 = df.head(10)
    last_10_avg_return = float(last_10["daily_return_pct"].mean(skipna=True))
    last_10_vol = float(last_10["daily_return_pct"].std(skipna=True))
    trend_pct = (float(last_10["close"].mean()) / float(first_10["close"].mean()) - 1) * 100
    metrics = {
        "symbol": str(df["symbol"].iloc[-1]) if "symbol" in df.columns else None,
        "n_rows": int(len(df)),
        "date_start": str(df["date"].iloc[0]),
        "date_end": str(df["date"].iloc[-1]),
        "close_start": round(start_close, 4),
        "close_end": round(end_close, 4),
        "net_change_pct": round(net_change_pct, 3),
        "close_high": round(high_close, 4),
        "close_low": round(low_close, 4),
        "avg_volume": round(avg_volume, 2) if avg_volume is not None else None,
        "last_10_avg_daily_return_pct": round(last_10_avg_return, 4),
        "last_10_volatility_pct": round(last_10_vol, 4),
        "trend_first10_to_last10_pct": round(trend_pct, 3),
    }
    return df, metrics


# 3. Format for AI #################################

def build_ai_payload(df_clean: pd.DataFrame, metrics: dict, max_rows: int = 20) -> dict:
    """Build a compact JSON payload for the LLM (metrics + recent rows)."""
    recent = df_clean.tail(max_rows).assign(date=lambda x: x["date"].astype(str))
    return {
        "report_type": "stock_eod_summary",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "metrics": metrics,
        "recent_rows": recent.to_dict(orient="records"),
        "notes": {"daily_return_pct": "percent", "volume": "shares"},
    }


# 4. Journalistic prompt + Ollama #################################

def check_ollama_running() -> bool:
    try:
        return requests.get(f"{OLLAMA_HOST}/api/tags", timeout=5).status_code == 200
    except requests.exceptions.RequestException:
        return False


def ollama_generate(prompt: str, model: str) -> str:
    url = f"{OLLAMA_HOST}/api/generate"
    body = {"model": model, "prompt": prompt, "stream": False}
    response = requests.post(url, json=body, timeout=120)
    response.raise_for_status()
    return response.json().get("response", "")


def build_prompt(payload_json: str, symbol: str, metrics: dict) -> str:
    """Prompt: company intro + mixed formatting + 100% quantified. Injects date range so model uses newest data only."""
    date_start = metrics.get("date_start", "")
    date_end = metrics.get("date_end", "")
    return f"""You are a financial journalist. Write an insights report about {symbol} using ONLY the JSON data below.

CRITICAL — DATES: The report covers exactly {date_start} through {date_end}. Use ONLY these dates and the numbers in the JSON. The data below is the sole source for all figures and dates.

Structure your report EXACTLY like this (use these section titles and mix bullets with short paragraphs):

# [{symbol}]

## About the company
- Write 2–4 short bullets: what this company does, what it's known for (e.g. products, sector). Keep it factual and brief. You may use well-known facts about the company; for {symbol} that is Apple Inc.

## The numbers
- Use bullets to list the key figures from the data: period (date_start to date_end), number of trading days (n_rows), price range (close_low to close_high), close_start → close_end, net_change_pct, avg_volume. Cite every number.
- Add 1–2 short sentences that interpret what those numbers mean (e.g. "That puts the stock down X% over the period.").

## Trends and volatility
- Bullets: reference last_10_avg_daily_return_pct, last_10_volatility_pct, trend_first10_to_last10_pct. Use the exact values.
- One short paragraph: what this suggests for near-term behavior (e.g. "Recent volatility of X% points to …").

## Bottom line
- One crisp paragraph or 2 bullets: the main takeaway for an investor or reader. Use at least one specific number from the data.

Rules:
- 100% quantified: every claim must cite a number from the JSON. No vague words without a figure.
- Be creative and readable: vary bullets and short paragraphs; use clear subheadings.
- Use only data from the JSON (and standard company facts for the intro). 
- Total length: about 250–350 words.

DATA (JSON):
{payload_json}
"""


# 5. Save as Word only #################################

def save_report_docx(report_text: str, path: str) -> None:
    """Write the AI report to Word. Handles # / ## / ### headings and bullets."""
    doc = Document()
    for line in report_text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(path)


# 6. Run #################################

print("\n🚀 API → process → AI report (Word only)...\n")
print(f"Symbol: {SYMBOL} | Days: {N_DAYS} | Model: {OLLAMA_MODEL}")

df_raw = query_marketstack_eod(symbol=SYMBOL, n_days=N_DAYS, api_key=MARKETSTACK_API_KEY)
df_clean, metrics = process_eod_data(df_raw)
payload = build_ai_payload(df_clean, metrics, max_rows=20)
payload_json = json.dumps(payload, indent=2)
prompt = build_prompt(payload_json, symbol=SYMBOL, metrics=metrics)

if not check_ollama_running():
    print("\n⚠️  Ollama not reachable at", OLLAMA_HOST)
    print("Start with: ollama serve   then  ollama pull smollm2:1.7b")
    raise SystemExit(1)

print("🤖 Generating report with Ollama...")
report = ollama_generate(prompt, model=OLLAMA_MODEL)
save_report_docx(report, OUTPUT_DOCX)
print("✅ Saved:", OUTPUT_DOCX, "\n")
